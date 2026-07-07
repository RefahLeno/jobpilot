import json
import os
import sys

from docx import Document


def add_heading(document, text, level=1):
    if text:
        document.add_heading(text, level=level)


def add_bullets(document, items):
    for item in items or []:
        if item:
            document.add_paragraph(str(item), style="List Bullet")


def add_numbered(document, items):
    for item in items or []:
        if item:
            document.add_paragraph(str(item), style="List Number")


def build_doc(variant, output_path):
    doc = Document()
    title = variant.get("draftContent", {}).get("title") or variant.get("name") or "求职简历草稿"
    positioning = variant.get("positioning", "")
    summary = variant.get("draftContent", {}).get("summary", "")
    skills = variant.get("draftContent", {}).get("skills", [])
    bullets = variant.get("draftContent", {}).get("bullets", [])
    rewrite_plan = variant.get("rewritePlan", [])
    warnings = variant.get("truthCheckWarnings", [])
    suitable_jds = variant.get("draftContent", {}).get("suitableJDs", [])

    add_heading(doc, title, level=0)
    if positioning:
        doc.add_paragraph(positioning)

    add_heading(doc, "个人摘要", level=1)
    doc.add_paragraph(summary or "待补充")

    add_heading(doc, "核心关键词", level=1)
    doc.add_paragraph("、".join(skills) if skills else "待补充")

    add_heading(doc, "改写重点", level=1)
    add_numbered(doc, rewrite_plan or ["待补充"])

    add_heading(doc, "建议 Bullet 草稿", level=1)
    add_bullets(doc, bullets or ["待补充"])

    add_heading(doc, "适合投递的岗位方向", level=1)
    add_bullets(doc, suitable_jds or ["待补充"])

    add_heading(doc, "真实性确认提醒", level=1)
    add_bullets(doc, warnings or ["请在导出前确认所有表述真实准确。"])

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def main():
    if len(sys.argv) < 3:
        print(json.dumps({"error": "missing arguments"}, ensure_ascii=False))
        sys.exit(1)

    payload_path = sys.argv[1]
    output_path = sys.argv[2]
    try:
        with open(payload_path, "r", encoding="utf-8-sig") as fh:
            payload = json.load(fh)
        variant = payload.get("variant") or {}
        build_doc(variant, output_path)
        print(json.dumps({"outputPath": output_path}, ensure_ascii=False))
    except Exception as exc:
        print(json.dumps({"error": str(exc)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
