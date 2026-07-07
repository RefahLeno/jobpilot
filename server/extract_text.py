import json
import os
import sys

from docx import Document
from pdfminer.high_level import extract_text as extract_pdf_text


def extract_docx(path):
    document = Document(path)
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "missing file path"}))
        sys.exit(1)

    path = sys.argv[1]
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            text = extract_pdf_text(path) or ""
        elif ext in [".docx", ".doc"]:
            text = extract_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        print(json.dumps({"text": text}, ensure_ascii=False))
    except Exception as exc:
        print(json.dumps({"error": str(exc)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
