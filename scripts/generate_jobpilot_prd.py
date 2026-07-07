from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT_DIR.mkdir(exist_ok=True)

BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY_FILL = "F2F4F7"


FULL_PATHS = [
    OUTPUT_DIR / "JobPilot-PRD-2026-07-06.docx",
    OUTPUT_DIR / "JobPilot-PRD-内测版-2026-07-06.docx",
    OUTPUT_DIR / "JobPilot-PRD-完整版-黑色版-2026-07-06.docx",
]
BRIEF_PATH = OUTPUT_DIR / "JobPilot-PRD-汇报版-面试版-黑色版-2026-07-06.docx"


def set_font(run, size=11, bold=False, color=BLACK, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def style_paragraph(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill=GRAY_FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK

    for style_name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    if "Body Small" not in doc.styles:
        body_small = doc.styles.add_style("Body Small", WD_STYLE_TYPE.PARAGRAPH)
        body_small.base_style = doc.styles["Normal"]
        body_small.font.name = "Calibri"
        body_small._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        body_small._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        body_small._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        body_small.font.size = Pt(10)
        body_small.font.color.rgb = BLACK
        body_small.paragraph_format.space_after = Pt(4)
        body_small.paragraph_format.line_spacing = 1.1


def add_title_block(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, before=0, after=4, line=1.0)
    set_font(p.add_run(title), size=21, bold=True)

    sub = doc.add_paragraph(style="Body Small")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(sub, before=0, after=14, line=1.1)
    set_font(sub.add_run(subtitle), size=10)


def add_paragraph(doc, text):
    p = doc.add_paragraph(style="Normal")
    style_paragraph(p, after=6, line=1.1)
    set_font(p.add_run(text))


def add_heading(doc, level, text):
    doc.add_paragraph(text, style=f"Heading {level}")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        set_font(p.add_run("• "), size=11)
        set_font(p.add_run(item), size=11)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        set_font(p.add_run(f"{index}. "), size=11, bold=True)
        set_font(p.add_run(item), size=11)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0)
        set_font(p.add_run(header), size=10.5, bold=True)

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            style_paragraph(p, after=0)
            set_font(p.add_run(value), size=10.5)

    doc.add_paragraph()


def build_meta_table(doc, rows):
    add_table(doc, ["字段", "内容"], rows, [Inches(1.8), Inches(4.7)])


def build_full_doc():
    doc = Document()
    configure_document(doc)
    add_title_block(
        doc,
        "JobPilot 求职工作台产品需求文档（PRD）",
        "文档版本：V1.1 | 文档日期：2026-07-06 | 文档类型：完整版黑色版",
    )
    build_meta_table(
        doc,
        [
            ("产品名称", "JobPilot 求职工作台"),
            ("产品定位", "面向求职者的简历与 JD 智能匹配工作台"),
            ("当前阶段", "内测版，已具备单 JD 分析、海投优化、登录、历史记录与 Word 导出主流程"),
            ("文档目标", "沉淀产品背景、功能范围、核心流程、阶段成果与后续规划"),
        ],
    )

    add_heading(doc, 1, "1. 项目背景")
    add_paragraph(doc, "很多求职者并不缺简历文件，而是缺少对岗位要求的理解、对自身匹配度的判断，以及针对不同岗位快速改写简历的能力。尤其在海投场景下，同一位用户往往要围绕多个岗位方向准备多版简历，决策成本和修改成本都很高。")
    add_paragraph(doc, "JobPilot 的目标不是单纯生成简历，而是帮助用户完成“理解 JD - 判断是否值得投递 - 找到差距 - 形成修改策略 - 生成可用版本”的完整求职决策流程。")

    add_heading(doc, 1, "2. 产品定位")
    add_bullets(doc, [
        "定位：面向求职者的简历与岗位 JD 智能匹配工作台。",
        "阶段：当前为内测版，优先服务真实用户试用与产品迭代，不以商业化收费为第一目标。",
        "价值：帮助用户更快看懂岗位、筛选值得投递的机会，并针对不同方向准备更贴合的简历版本。",
    ])

    add_heading(doc, 1, "3. 目标用户")
    add_bullets(doc, [
        "应届生与实习求职者：缺少岗位判断经验，不知道如何针对 JD 调整简历。",
        "转行或跨方向求职者：需要评估自己与目标岗位之间的能力差距和表达差距。",
        "海投求职者：需要将大量 JD 归类成若干投递方向，并高效准备多版本简历。",
        "职业辅导或内推顾问：需要借助结构化结果快速给出简历与投递建议。",
    ])

    add_heading(doc, 1, "4. 核心问题定义")
    add_bullets(doc, [
        "用户看不懂 JD 中的真实能力要求，无法判断岗位是否值得重点投递。",
        "用户无法清晰识别“我已经匹配什么”和“我还差什么”，只能凭感觉改简历。",
        "海投场景下，用户面对大量相似但不完全相同的 JD，难以判断应准备几版简历。",
        "很多工具停留在“生成内容”，但缺少对真实性、可解释性和可执行性的约束。",
    ])

    add_heading(doc, 1, "5. 产品目标")
    add_bullets(doc, [
        "让用户在 5 分钟内完成一条单 JD 分析链路，并获得可执行的修改建议。",
        "让用户在 15 分钟内完成一批 JD 的分类与多版本简历策略生成。",
        "让报告输出不只停留在打分，而是能明确指出匹配点、缺口、优先修改项和下一步动作。",
        "保证所有生成结果遵守“保守增强”原则：不编造经历，不新增未提供事实，只做表达重组与建议生成。",
    ])

    add_heading(doc, 1, "6. 当前产品范围")
    add_heading(doc, 2, "6.1 已覆盖能力")
    add_bullets(doc, [
        "简历上传与解析：支持 PDF / DOC / DOCX 上传、文本提取、PDF 预览、Word 文本预览。",
        "简历结构化结果：展示摘要、关键词和分组关键词。",
        "单 JD 分析：支持手动粘贴 JD 或通过链接抓取 JD，并生成匹配报告。",
        "海投优化：支持批量导入 JD 文本与链接，进行岗位方向分类，并生成多版本简历草稿。",
        "账号与历史：支持邮箱注册登录、历史记录回看、多份简历管理、版本详情查看。",
        "导出能力：支持结构化 Word 导出。",
    ])
    add_heading(doc, 2, "6.2 暂不纳入范围")
    add_bullets(doc, [
        "不做本地大模型部署与 GPU 推理。",
        "不做扫描版 PDF OCR。",
        "不做复杂 ATS 模板系统与不可控的一键成稿。",
        "不做支付、套餐、短信验证、团队协作等商业化能力。",
    ])

    add_heading(doc, 1, "7. 核心流程设计")
    add_heading(doc, 2, "7.1 单 JD 分析流程")
    add_numbered(doc, [
        "用户上传基础简历，系统完成文本提取、关键词抽取和预览展示。",
        "用户输入 JD 文本或粘贴 JD 链接，系统抓取并标准化 JD 内容。",
        "系统对简历与 JD 分别进行总结、关键词提取与匹配分析。",
        "系统输出匹配分、维度评分、匹配点、缺失点、优先修改方向和下一步建议。",
        "用户根据报告优化简历，并决定是否投递该岗位。",
    ])
    add_heading(doc, 2, "7.2 海投优化流程")
    add_numbered(doc, [
        "用户上传基础简历并进入海投模式。",
        "用户批量导入多个 JD 文本或链接。",
        "系统解析每个 JD 的标题、关键词、方向标签和基础匹配度。",
        "系统将 JD 聚类为若干岗位方向。",
        "系统为每个方向生成简历修改策略、多版本草稿和适配 JD 列表。",
        "用户保存、回看并导出目标版本，用于后续投递。",
    ])

    add_heading(doc, 1, "8. 功能模块设计")
    add_table(
        doc,
        ["模块", "功能说明"],
        [
            ("简历中心", "上传多份简历、查看摘要、关键词、上传时间与详情，选择当前分析所使用的基础简历。"),
            ("单 JD 工作台", "围绕一个目标岗位完成匹配分析，重点回答“值不值得投、差在哪里、先改什么”。"),
            ("海投优化工作台", "围绕一批 JD 完成分类与多版本策略生成，重点回答“要准备几版简历、每版覆盖哪些岗位”。"),
            ("历史记录", "回看最近的单 JD 报告、海投批次和简历版本，支撑连续使用。"),
            ("版本管理", "保存版本命名、投递状态、主推标记和导出记录。"),
            ("管理员监控", "仅管理员可见，用于查看请求、错误和系统运行概况。"),
        ],
        [Inches(1.7), Inches(4.8)],
    )

    add_heading(doc, 1, "9. 报告与结果设计原则")
    add_bullets(doc, [
        "报告第一优先级不是“评分”，而是“指导用户先改哪里”。",
        "结果结构按“先改什么 - 为什么 - 已匹配什么 - 具体怎么改”组织。",
        "将“能力差距”“信息缺口”“投递前确认项”区分展示，避免误判。",
        "整体保持中文化、可解释、可执行的输出风格。",
    ])

    add_heading(doc, 1, "10. 海投优化设计原则")
    add_bullets(doc, [
        "批量 JD 的目标不是逐条打分，而是帮助用户形成岗位方向分组和简历版本策略。",
        "岗位分类结果必须对求职者可理解，例如 AI 产品经理、内容产品经理、增长产品经理等方向。",
        "每个简历版本都需说明强化哪些关键词、弱化哪些内容、适合覆盖哪些 JD、有哪些真实性风险需要确认。",
        "第一版只生成草稿和建议，不直接生成不可控的最终简历成品。",
    ])

    add_heading(doc, 1, "11. AI 与规则策略")
    add_bullets(doc, [
        "DeepSeek 负责结构化总结、关键词提取、结果解释和简历改写建议。",
        "本地规则与启发式分析作为 fallback，保证 API 不可用时主流程仍可运行。",
        "批量 JD 分类同时依赖关键词、职责文本和语义相似度，不只做纯关键词匹配。",
        "所有生成结果都遵守保守增强原则，不虚构经历和成果。",
    ])

    add_heading(doc, 1, "12. 数据与权限")
    add_bullets(doc, [
        "当前采用 Node 单服务 + 本地 JSON 持久化架构，适合内测与轻量试用阶段。",
        "所有简历、JD、报告、批次、版本记录都带有用户归属信息，仅当前登录用户可访问。",
        "管理员监控页与相关接口仅对白名单管理员开放，不在普通用户主页暴露。",
    ])

    add_heading(doc, 1, "13. 指标建议")
    add_table(
        doc,
        ["指标", "建议口径"],
        [
            ("单 JD 完成率", "上传简历后成功进入分析结果页的用户占比。"),
            ("海投批次完成率", "成功完成“导入 JD - 聚类 - 生成版本”的批次占比。"),
            ("多版本产出率", "每个海投批次最终生成 2 到 5 个有效版本的占比。"),
            ("报告可执行性反馈", "用户反馈中“建议具体、能指导修改”的正向评价比例。"),
            ("留存与复用", "用户再次回到历史记录、再次使用已上传简历或已生成版本的比例。"),
        ],
        [Inches(1.8), Inches(4.7)],
    )

    add_heading(doc, 1, "14. 里程碑规划")
    add_table(
        doc,
        ["阶段", "目标与说明"],
        [
            ("阶段一：稳定 MVP", "打通单 JD 和海投主流程，补齐异常状态、加载状态和基本导出能力。"),
            ("阶段二：AI 与持久化", "正式接入 DeepSeek，补齐 provider / fallback 可观测性，并保证结果可保存、可回看。"),
            ("阶段三：连续使用能力", "完善多份简历管理、历史批次详情、版本管理与导出复用。"),
            ("阶段四：上线基础", "完成账号体系、数据隔离、调用限制、日志监控与管理员后台。"),
        ],
        [Inches(1.7), Inches(4.8)],
    )

    add_heading(doc, 1, "15. 风险与限制")
    add_bullets(doc, [
        "JD 链接抓取可能受到登录态、反爬或动态渲染影响，因此必须保留手动粘贴兜底。",
        "扫描版 PDF 暂不支持 OCR，会影响部分简历解析成功率。",
        "本地 JSON 存储适合内测，但不适合长期高并发生产使用，后续需迁移数据库。",
        "AI 输出需持续通过 bad case 迭代，尤其要避免误判信息缺口和能力缺口。",
    ])

    add_heading(doc, 1, "16. 验收标准")
    add_bullets(doc, [
        "用户可以完成注册登录，并在刷新后保留登录态。",
        "用户可以上传 PDF / Word 简历，获得预览、关键词和摘要结果。",
        "用户可以完成单 JD 分析，并获取全中文、可解释、可执行的匹配报告。",
        "用户可以批量导入 5 到 30 条 JD，完成岗位方向分类与多版本策略生成。",
        "用户可以回看历史报告、历史海投批次和简历版本，并完成 Word 导出。",
    ])

    add_heading(doc, 1, "17. 后续迭代建议")
    add_bullets(doc, [
        "优化报告判定规则，进一步区分能力差距、信息缺口与投递确认项。",
        "增强海投分类命名和版本差异化质量，让生成结果更贴近真实岗位族群。",
        "逐步迁移至数据库与更稳定的部署架构，支撑真实内测用户使用。",
        "在商业化阶段接入支付、额度策略和更完整的用户运营机制。",
    ])
    return doc


def build_brief_doc():
    doc = Document()
    configure_document(doc)
    add_title_block(
        doc,
        "JobPilot 求职工作台汇报版 / 面试版 PRD",
        "文档版本：V1.0 | 文档日期：2026-07-06 | 文档类型：压缩汇报版黑色版",
    )
    build_meta_table(
        doc,
        [
            ("项目名称", "JobPilot 求职工作台"),
            ("适用场景", "项目汇报、面试介绍、简历项目补充说明"),
            ("核心目标", "用更短的文档讲清产品背景、功能链路、亮点和阶段成果"),
            ("阅读方式", "适合 3 到 5 分钟快速浏览，也适合口头讲述时对照使用"),
        ],
    )

    add_heading(doc, 1, "1. 项目一句话概述")
    add_paragraph(doc, "JobPilot 是一个面向求职者的 AI 求职工作台，核心价值是帮助用户看懂 JD、判断匹配度、找到简历差距，并在海投场景下快速形成多版本简历策略。")

    add_heading(doc, 1, "2. 为什么做这个项目")
    add_bullets(doc, [
        "求职者常常看不懂 JD 的真实要求，不知道岗位值不值得投。",
        "很多人知道自己要改简历，但不知道该先改哪里、怎么改。",
        "海投场景下，同一位用户往往要面对很多相似但不完全相同的岗位，简历改版成本很高。",
    ])

    add_heading(doc, 1, "3. 产品定位")
    add_bullets(doc, [
        "不是单纯的简历生成工具，而是求职决策辅助工作台。",
        "第一条主线是单 JD 分析，解决“这个岗位值不值得投”的问题。",
        "第二条主线是海投优化，解决“这批岗位该分成哪几类、要准备几版简历”的问题。",
    ])

    add_heading(doc, 1, "4. 核心用户")
    add_bullets(doc, [
        "应届生和实习求职者",
        "转行或跨方向求职者",
        "海投用户",
        "职业辅导或内推顾问",
    ])

    add_heading(doc, 1, "5. 核心功能")
    add_table(
        doc,
        ["功能", "说明"],
        [
            ("简历解析", "支持 PDF / Word 上传、文本提取、摘要与关键词展示。"),
            ("单 JD 分析", "输入 JD 后生成匹配分、匹配点、差距点和优化建议。"),
            ("海投优化", "批量导入 JD，自动完成岗位方向分类与多版本简历草稿生成。"),
            ("历史与版本", "支持历史记录回看、多份简历管理、版本保存和 Word 导出。"),
        ],
        [Inches(1.8), Inches(4.7)],
    )

    add_heading(doc, 1, "6. 核心流程")
    add_heading(doc, 2, "6.1 单 JD")
    add_numbered(doc, [
        "上传简历",
        "输入或抓取 JD",
        "生成匹配报告",
        "根据建议修改简历",
        "决定是否投递",
    ])
    add_heading(doc, 2, "6.2 海投优化")
    add_numbered(doc, [
        "上传基础简历",
        "批量导入多个 JD",
        "系统完成岗位方向分类",
        "生成不同方向的简历修改策略与草稿",
        "保存、回看并导出版本",
    ])

    add_heading(doc, 1, "7. 关键产品设计思路")
    add_bullets(doc, [
        "把单 JD 和海投拆成两个模式，因为两者的用户任务目标不同。",
        "报告重心不是打分本身，而是告诉用户“先改什么”。",
        "海投场景强调“岗位方向分类”和“多版本策略”，而不是单纯批量跑分。",
        "所有生成结果都遵守保守增强原则，不编造经历，只做表达重组和策略建议。",
    ])

    add_heading(doc, 1, "8. 当前阶段成果")
    add_bullets(doc, [
        "已完成前端工作台和主要后端接口。",
        "已支持注册登录、数据隔离、历史记录、多份简历管理和版本管理。",
        "已支持单 JD 分析、海投优化、Word 导出与管理员监控页。",
        "已具备内测版可运行闭环，可以支撑真实用户试用和继续迭代。",
    ])

    add_heading(doc, 1, "9. 项目亮点")
    add_bullets(doc, [
        "从真实求职痛点出发，不是为了做 AI 而做 AI。",
        "既考虑用户体验，也考虑结果真实性和解释性。",
        "支持从单次分析走向连续使用，具备工作台属性。",
        "在海投场景下引入岗位分类与多版本策略，更贴近真实求职行为。",
    ])

    add_heading(doc, 1, "10. 后续规划")
    add_bullets(doc, [
        "继续优化报告规则，区分能力差距、信息缺口和投递确认项。",
        "增强海投分类质量和版本差异化质量。",
        "逐步完善部署和持久化能力，支撑更稳定的真实用户使用。",
        "后续再考虑支付、额度策略和商业化能力。",
    ])
    return doc


def save_all():
    full_doc = build_full_doc()
    for path in FULL_PATHS:
        full_doc.save(path)

    brief_doc = build_brief_doc()
    brief_doc.save(BRIEF_PATH)

    print(str(FULL_PATHS[-1]))
    print(str(BRIEF_PATH))


if __name__ == "__main__":
    save_all()
